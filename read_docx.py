#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

doc = Document(r'c:\Shahin-ai\Migration_and_Onboarding_Verification_then_Action_Playbook.docx')

print("=" * 80)
print("MIGRATION AND ONBOARDING VERIFICATION PLAYBOOK")
print("=" * 80)
print()

for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

print()
for i, table in enumerate(doc.tables):
    print(f"\n{'=' * 40}")
    print(f"TABLE {i+1}")
    print("=" * 40)
    for row in table.rows:
        cells = [cell.text.strip().replace('\n', ' ')[:50] for cell in row.cells]
        print(" | ".join(cells))
